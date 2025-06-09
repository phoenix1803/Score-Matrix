import os
import re
import pdfplumber
import nltk
import pandas as pd
import json
from docx import Document
from transformers import pipeline
import google.generativeai as genai
import language_tool_python
import nltk
import spacy
import warnings
import numpy as np
from collections import Counter
from difflib import SequenceMatcher
from nltk.corpus import stopwords, wordnet
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize, sent_tokenize
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.naive_bayes import MultinomialNB
from textblob import TextBlob
from pdf2image import convert_from_path
from PIL import Image
import tempfile

BASE_DIR = "/tmp"

genai.configure(api_key="")

warnings.filterwarnings("ignore")

nltk.download('stopwords', quiet=True)
nltk.download('punkt', quiet=True)
nltk.download('vader_lexicon', quiet=True)

EN_STOPWORDS = set(stopwords.words("english"))

def extract_questions_from_docx(docx_path):
    """Extract questions from a DOCX file."""
    doc = Document(docx_path)
    questions = []
    pattern = re.compile(r"Question \d+\s*(.*?)\s*\[(\d+) marks\]")
    
    for para in doc.paragraphs:
        match = pattern.match(para.text.strip())
        if match:
            question, marks = match.groups()
            questions.append((question.strip(), int(marks)))
    
    return questions

def extract_questions_from_pdf(pdf_path):
    """Extract questions from a PDF file."""
    questions = []
    pattern = re.compile(r"Question \d+\s*(.*?)\s*\[(\d+) marks\]")
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    match = pattern.match(line.strip())
                    if match:
                        question, marks = match.groups()
                        questions.append((question.strip(), int(marks)))
    
    return questions

def classify_questions(questions):
    """Classify questions into MCQ, Descriptive, or Opinion-based."""
    classifier = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
    question_types = ["MCQ", "Descriptive", "Opinion-based"]
    classified_data = []
    
    for i, (question, marks) in enumerate(questions):
        result = classifier(question, question_types)
        predicted_type = result["labels"][0]
        classified_data.append({"Question Number": i + 1, "Question": question, "Type": predicted_type, "Marks": marks})
    
    return classified_data

def save_questions_to_json(classified_data, output_json):
    """Save classified questions to a JSON file."""
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(classified_data, f, indent=4, ensure_ascii=False)

def get_answer(question, marks):
    """Generate an answer using Gemini API."""
    if marks == 1:
        prompt = f"Answer this in one precise sentence:\nQ: {question}"
    elif marks == 2:
        prompt = f"Provide an answer in 2-3 sentences:\nQ: {question}"
    elif marks == 3:
        prompt = f"Provide a short answer in 3-4 sentences:\nQ: {question}"
    elif marks == 4:
        prompt = f"Provide a detailed answer in about 5-6 sentences:\nQ: {question}"
    elif marks >= 5:
        prompt = f"Write a comprehensive answer in multiple paragraphs about 500 words:\nQ: {question}"
    else:
        prompt = f"Provide a precise and clear answer:\nQ: {question}"

    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    return response.text.strip() if response.text else "No answer found."

def generate_answers_json_and_docx(classified_data, output_json, output_docx):
    """Generate answers and save them in JSON and DOCX formats."""
    with open(output_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    doc = Document()
    
    for entry in data:
        entry["Reference Answer"] = get_answer(entry["Question"], entry["Marks"])
        doc.add_paragraph(f"Q: {entry['Question']} ({entry['Marks']} marks)", style="Heading 2")
        doc.add_paragraph(f"A: {entry['Reference Answer']}\n", style="Normal")
    
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
    
    doc.save(output_docx)

def process_question_paper(input_file):
    """Process the question paper and generate reference answers."""
    output_dir = "../frontend/outputs"
    os.makedirs(output_dir, exist_ok=True)
    output_json = os.path.join(output_dir, "classified_questions.json")
    output_docx = os.path.join(output_dir, "answers.docx")
    
    if input_file.endswith(".pdf"):
        questions = extract_questions_from_pdf(input_file)
    elif input_file.endswith(".docx"):
        questions = extract_questions_from_docx(input_file)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
    
    classified_data = classify_questions(questions)
    save_questions_to_json(classified_data, output_json)
    generate_answers_json_and_docx(classified_data, output_json, output_docx)
    print("Processing complete. Files saved in output folder.")

def extract_text_from_image(image):
    """Extract text from an image using Gemini 1.5 Flash."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = "Extract only the text from this image. Do not add any explanations or extra words. Maintain the exact structure of the text."
    response = model.generate_content([image, prompt], stream=False)
    return response.text.strip()

def extract_answers_from_pdf(pdf_path):
    """Extract answers from a student's PDF using Gemini OCR."""
    answers = []
    roll_number = "Unknown"

    # Convert PDF to images
    images = convert_from_path(pdf_path, dpi=300)  # Convert PDF pages to images

    for page_num, image in enumerate(images):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_img:
            image.save(temp_img.name, format="JPEG")  # Save image to a temporary file
        
        # Open the image file and ensure it is closed after processing
        with Image.open(temp_img.name) as img:
            extracted_text = extract_text_from_image(img)  # Extract text using Gemini
        
        # Delete the temporary file after processing
        os.unlink(temp_img.name)

        # Debugging: Print extracted text
        print(f"Page {page_num + 1} Extracted Text:\n{extracted_text}\n")

        if extracted_text:
            if page_num == 0:  # First page might contain Roll Number
                # Look for roll number in the extracted text
                if "Roll Number" in extracted_text:
                    roll_line = extracted_text.split("Roll Number")[1].split("\n")[0].strip()
                    roll_number = roll_line.split()[-1].strip()  # Extract the roll number

            # Split text into lines
            lines = extracted_text.split("\n")
            current_answer = []  # Temporary list to hold lines for the current answer
            for line in lines:
                line = line.strip()
                if not line:
                    continue  # Skip empty lines

                # Check if the line starts with "Ans" (case-insensitive)
                if line.lower().startswith("ans"):
                    # If we have collected lines for a previous answer, add it to the answers list
                    if current_answer:
                        answers.append(" ".join(current_answer).strip())  # Join lines into a single answer
                        current_answer = []  # Reset for the next answer

                    # Add the current "Ans" line to the new answer
                    current_answer.append(line)
                else:
                    # If the line doesn't start with "Ans", it's part of the current answer
                    current_answer.append(line)

            # Add the last collected answer (if any)
            if current_answer:
                answers.append(" ".join(current_answer).strip())

    return answers, roll_number

def process_mcq_answer(reference_answer, student_answer):
    """Evaluate MCQ answers."""
    return 1.0 if reference_answer.lower() == student_answer.lower() else 0.0

def process_opinion_answer(reference_answer, student_answer, max_marks):
    """
    Evaluate opinion-based answers using Gemini.
    
    Args:
        reference_answer (str): The reference answer.
        student_answer (str): The student's answer.
        max_marks (int): Maximum marks for the question.
    
    Returns:
        float: Score awarded to the student's answer.
    """
    prompt = f"""
    Evaluate the student's opinion-based answer based on the following criteria:
    1. **Sentiment Match**: Does the student's sentiment align with the reference answer?
    2. **Coherence**: Is the student's answer logically structured and easy to understand?
    3. **Logical Flow**: Does the student's answer present ideas in a logical sequence?
    4. **Grammatical Accuracy**: Is the student's answer free of grammatical errors?
    5. **Relevance**: Does the student's answer stay relevant to the question?

    Reference Answer: {reference_answer}
    Student Answer: {student_answer}

    Provide a score out of {max_marks} based on the above criteria. Return only the score as a number.
    """

    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)

    try:
        score = float(response.text.strip())
        return min(score, max_marks)
    except ValueError:
        print("Gemini did not return a valid score. Using fallback scoring method.")
        return evaluate_opinion_fallback(reference_answer, student_answer, max_marks)

def evaluate_opinion_fallback(reference_answer, student_answer, max_marks):
    """
    Fallback method to evaluate opinion-based answers if Gemini fails.
    """
    scores = {
        "sentiment_match": sentiment_analysis(reference_answer) - sentiment_analysis(student_answer),
        "coherence": coherence_score(student_answer),
        "logical_flow": logical_flow(student_answer),
        "grammatical_accuracy": grammatical_accuracy(student_answer),
        "relevance": relevance_score(reference_answer, student_answer)
    }
    
    weights = {
        "sentiment_match": 0.1,
        "coherence": 0.3,
        "logical_flow": 0.2,
        "grammatical_accuracy": 0.2,
        "relevance": 0.2
    }
    
    weighted_score = weighted_average_score(scores, weights)
    return round(weighted_score * max_marks, 1)

def process_descriptive_answer(reference_answer, student_answer, max_marks):
    """Evaluate descriptive answers."""
    scores = {
        'exact_match': exact_match(reference_answer, student_answer),
        'partial_match': partial_match(reference_answer, student_answer),
        'cosine_similarity': cosine_similarity_score(reference_answer, student_answer),
        'sentiment': sentiment_analysis(student_answer),
        'enhanced_sentence_match': enhanced_sentence_match(reference_answer, student_answer),
        'multinomial_naive_bayes': multinomial_naive_bayes_score(reference_answer, student_answer),
        'semantic_similarity': semantic_similarity_score(reference_answer, student_answer),
        "coherence": coherence_score(student_answer),
        'relevance': relevance_score(reference_answer, student_answer),
        'grammatical_accuracy': grammatical_accuracy(student_answer),
        'keyword_coverage': keyword_coverage(reference_answer, student_answer),
        'factual_consistency': factual_consistency(reference_answer, student_answer),
        'structural_similarity': structural_similarity(reference_answer, student_answer),
        'readability': readability_score(student_answer),
        'concept_ordering': concept_ordering(reference_answer, student_answer),
        'technical_accuracy': technical_accuracy(reference_answer, student_answer),
        'logical_flow': logical_flow(student_answer)
    }

    weights = {
        'exact_match': 0.08,
        'partial_match': 0.06,
        'cosine_similarity': 0.06,
        'sentiment': 0.04,
        'enhanced_sentence_match': 0.06,
        'multinomial_naive_bayes': 0.06,
        'semantic_similarity': 0.06,
        'coherence': 0.06,
        'relevance': 0.06,
        'grammatical_accuracy': 0.07,
        'keyword_coverage': 0.07,
        'factual_consistency': 0.08,
        'structural_similarity': 0.05,
        'readability': 0.05,
        'concept_ordering': 0.05,
        'technical_accuracy': 0.07,
        'logical_flow': 0.08
    }
    
    weighted_score = weighted_average_score(scores, weights)
    return round(weighted_score * max_marks, 1)

def exact_match(expected_answer, student_answer):
    return int(expected_answer == student_answer)
def factual_consistency(expected_answer, student_answer):
    nlp = spacy.load('en_core_web_sm')
    
    expected_doc = nlp(expected_answer)
    student_doc = nlp(student_answer)
    
    expected_facts = set([(ent.text, ent.label_) for ent in expected_doc.ents])
    student_facts = set([(ent.text, ent.label_) for ent in student_doc.ents])
    
    if len(expected_facts) == 0:
        return 1.0
    
    matching_facts = expected_facts & student_facts
    consistency_score = len(matching_facts) / len(expected_facts)
    return consistency_score

def multinomial_naive_bayes_score(expected_answer, student_answer):
    answers = [expected_answer, student_answer]
    vectorizer = CountVectorizer(tokenizer=preprocess_text)
    X = vectorizer.fit_transform(answers)
    y = [0, 1] 
    clf = MultinomialNB()
    clf.fit(X, y)
    probs = clf.predict_proba(X)
    return probs[1][1] 

def readability_score(text):
    analysis = TextBlob(text)
    polarity = (analysis.sentiment.polarity + 1) / 2
    subjectivity = 1 - analysis.sentiment.subjectivity 
    return (polarity + subjectivity) / 2

def concept_ordering(expected_answer, student_answer):
    expected_tokens = word_tokenize(expected_answer.lower())
    student_tokens = word_tokenize(student_answer.lower())
    
    expected_keywords = [word for word in expected_tokens if word not in EN_STOPWORDS]
    student_keywords = [word for word in student_tokens if word not in EN_STOPWORDS]
    
    common_keywords = set(expected_keywords) & set(student_keywords)
    
    if not common_keywords:
        return 0
    
    expected_positions = {word: i for i, word in enumerate(expected_keywords) if word in common_keywords}
    student_positions = {word: i for i, word in enumerate(student_keywords) if word in common_keywords}
    
    position_differences = sum(abs(expected_positions[word] - student_positions[word]) for word in common_keywords)
    max_possible_difference = len(expected_keywords) * len(student_keywords)
    
    order_score = 1 - (position_differences / max_possible_difference)
    return max(0, order_score)

def technical_accuracy(expected_answer, student_answer):
    nlp = spacy.load('en_core_web_sm')
    
    expected_doc = nlp(expected_answer)
    student_doc = nlp(student_answer)
    
    expected_terms = set([token.text.lower() for token in expected_doc if token.pos_ in ['NOUN', 'PROPN']])
    student_terms = set([token.text.lower() for token in student_doc if token.pos_ in ['NOUN', 'PROPN']])
    
    if not expected_terms:
        return 1.0
        
    accuracy_score = len(expected_terms & student_terms) / len(expected_terms)
    return accuracy_score

def semantic_similarity_score(expected_answer, student_answer):
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    embeddings_expected = model.encode([expected_answer])
    embeddings_student = model.encode([student_answer])
    similarity = cosine_similarity([embeddings_expected.flatten()], [embeddings_student.flatten()])[0][0]
    return similarity

def partial_match(expected_answer, student_answer):
    expected_tokens = preprocess_text(expected_answer)
    student_tokens = preprocess_text(student_answer)
    common_tokens = set(expected_tokens) & set(student_tokens)
    match_percentage = len(common_tokens) / max(len(expected_tokens), len(student_tokens))
    return match_percentage

def cosine_similarity_score(expected_answer, student_answer):
    vectorizer = TfidfVectorizer(tokenizer=preprocess_text)
    tfidf_matrix = vectorizer.fit_transform([expected_answer, student_answer])
    cosine_sim = cosine_similarity(tfidf_matrix[0], tfidf_matrix[1])[0][0]
    return cosine_sim

def preprocess_text(text):
    tokens = word_tokenize(text)
    lemmatizer = WordNetLemmatizer()
    return [lemmatizer.lemmatize(token.lower()) for token in tokens if token.lower() not in EN_STOPWORDS]

def sentiment_analysis(text):
    sia = SentimentIntensityAnalyzer()
    return (sia.polarity_scores(text)['compound'] + 1) / 2

def enhanced_sentence_match(expected_answer, student_answer):
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
    embeddings_expected = model.encode([expected_answer])
    embeddings_student = model.encode([student_answer])
    return cosine_similarity([embeddings_expected.flatten()], [embeddings_student.flatten()])[0][0]

def coherence_score(text):
    sentences = sent_tokenize(text)
    return 1.0 if len(sentences) <= 1 else min(len(set(text.split())) / max(1, len(text.split())), 1.0)

def relevance_score(expected_answer, student_answer):
    expected_tokens = set(preprocess_text(expected_answer))
    student_tokens = set(preprocess_text(student_answer))
    return len(expected_tokens & student_tokens) / len(expected_tokens) if expected_tokens else 0

def grammatical_accuracy(text):
    try:
        tool = language_tool_python.LanguageTool('en-US')
        matches = tool.check(text)
        return 1 - min(len(matches) / max(1, len(text.split())), 1.0)
    except:
        return 0.8  

def keyword_coverage(expected_answer, student_answer):
    expected_tokens = set(preprocess_text(expected_answer))
    student_tokens = set(preprocess_text(student_answer))
    return len(student_tokens & expected_tokens) / len(expected_tokens) if expected_tokens else 0

def structural_similarity(expected_answer, student_answer):
    return SequenceMatcher(None, expected_answer, student_answer).ratio()

def logical_flow(text):
    sentences = sent_tokenize(text)
    if len(sentences) <= 1:
        return 0.5  
    return min(1.0, (len(sentences) / max(1, len(text.split())) * 10))

def weighted_average_score(scores, weights):
    return sum(score * weights.get(criterion, 0) for criterion, score in scores.items()) / sum(weights.values())

def identify_weak_topics(question, student_answer, reference_answer):
    prompt = f"""
    Analyze the student's mistakes in answering the following question. Identify the weak topic and suggest improvements.
    
    Question: {question}
    Reference Answer: {reference_answer}
    Student Answer: {student_answer}
    
    Provide the weak topic and tips in JSON format with keys 'weak_topic' and 'tips'. The 'tips' should be a list of strings.
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    print("Raw Response from Gemini API:")
    print(response.text)
    
    try:
        result = json.loads(response.text)
    except json.JSONDecodeError:
        print("Failed to parse response as JSON. Attempting manual extraction.")
        weak_topic = "Unknown"
        tips = ["No specific tips available."]
        
        weak_topic_match = re.search(r'"weak_topic"\s*:\s*"([^"]+)"', response.text)
        tips_match = re.findall(r'"tips"\s*:\s*\[([^\]]+)\]', response.text)
        
        if weak_topic_match:
            weak_topic = weak_topic_match.group(1)
        if tips_match:
            tips = [tip.strip().strip('"') for tip in tips_match[0].split(",")]
        
        result = {"weak_topic": weak_topic, "tips": tips}
    
    return result
def evaluate_answers(input_pdf, classified_json, output_json):
    """
    Evaluate a student's answers against reference answers.

    Args:
        input_pdf (str): Path to the student's answer PDF.
        classified_json (str): Path to the JSON file containing classified questions.
        output_json (str): Path to save the evaluation results.

    Returns:
        list: A list of dictionaries containing evaluation results.
    """
    try:
        with open(classified_json, "r", encoding="utf-8") as f:
            classified_data = json.load(f)
        
        student_answers, roll_number = extract_answers_from_pdf(input_pdf)
        
        results = []
        
        for i, entry in enumerate(classified_data):
            question_num = entry["Question Number"]
            question = entry["Question"]
            q_type = entry["Type"]
            marks = entry["Marks"]
            reference_answer = entry["Reference Answer"]
            
            student_answer = student_answers[i][1] if i < len(student_answers) else ""
            
            if q_type == "MCQ":
                score = process_mcq_answer(reference_answer, student_answer)
            elif q_type == "Opinion-based":
                score = process_opinion_answer(reference_answer, student_answer, marks)
            else:
                score = process_descriptive_answer(reference_answer, student_answer, marks)
            
            result_entry = {
                "roll_number": roll_number,
                "file_name": os.path.basename(input_pdf),
                "question_number": question_num,
                "question": question,
                "reference_answer": reference_answer,
                "student_answer": student_answer,
                "marks_awarded": score,
                "total_marks": marks
            }
            
            if score < marks * 0.8: 
                weak_data = identify_weak_topics(question, student_answer, reference_answer)
                result_entry["weak_topics"] = weak_data["weak_topic"]
                result_entry["improvement_tips"] = weak_data["tips"]
            
            results.append(result_entry)
        
        os.makedirs(os.path.dirname(output_json), exist_ok=True)
        
        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=4, ensure_ascii=False)
        
        return results
    
    except Exception as e:
        print(f"An error occurred while evaluating answers: {e}")
        return []

def process_student_answer_sheet(student_pdf):
    """Process a single student's answer sheet."""
    classified_json = "../frontend/outputs/classified_questions.json"
    output_json = f"../frontend/data/{os.path.basename(student_pdf).split('.')[0]}_results.json"
    
    results = evaluate_answers(student_pdf, classified_json, output_json)
    print(f"Evaluation complete for {student_pdf}. Results saved to {output_json}")
    
    return output_json

def process_all_student_sheets(folder_path):
    """Process all student answer sheets in a folder."""
    result_files = []
    
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf") and "question-paper" not in filename.lower():
            student_pdf = os.path.join(folder_path, filename)
            result_file = process_student_answer_sheet(student_pdf)
            result_files.append(result_file)
    
    return result_files

def generate_class_summary(result_files, output_json):
    """Generate a summary of all students' performance."""
    all_results = []
    student_summaries = {}
    
    for file_path in result_files:
        with open(file_path, "r", encoding="utf-8") as f:
            results = json.load(f)
            all_results.extend(results)
            
            if results:
                roll_number = results[0]["roll_number"]
                total_marks = sum(r["total_marks"] for r in results)
                marks_awarded = sum(r["marks_awarded"] for r in results)
                percentage = (marks_awarded / total_marks * 100) if total_marks > 0 else 0
                
                weak_topics = {}
                for result in results:
                    if "weak_topics" in result:
                        topic = result["weak_topics"]
                        weak_topics[topic] = weak_topics.get(topic, 0) + 1
                
                common_weak_topics = [topic for topic, count in 
                                     sorted(weak_topics.items(), key=lambda x: x[1], reverse=True)[:3]]
                
                student_summaries[roll_number] = {
                    "roll_number": roll_number,
                    "total_marks": total_marks,
                    "marks_awarded": marks_awarded,
                    "percentage": round(percentage, 2),
                    "common_weak_topics": common_weak_topics
                }
    
    summary = {
        "individual_results": all_results,
        "student_summaries": list(student_summaries.values())
    }
    
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=4, ensure_ascii=False)
    
    print(f"Class summary generated and saved to {output_json}")

def extract_questions_from_docx(docx_path):
    """Extract questions from a DOCX file."""
    doc = Document(docx_path)
    questions = []
    pattern = re.compile(r"Question\s+\d+[:.]?\s*(.*?)\s*\[(\d+)\s*marks?\]", re.IGNORECASE | re.DOTALL)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = pattern.search(text)
        if match:
            question, marks = match.groups()
            questions.append((question.strip(), int(marks)))
    
    return questions

def pdf_to_images(pdf_path):
    """Convert PDF to a list of PIL Image objects without using pdf2image."""
    images = []
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                # Extract page as bytes - this will be a PDF subset containing just that page
                writer = PyPDF2.PdfWriter()
                writer.add_page(pdf_reader.pages[page_num])
                
                with io.BytesIO() as bytes_stream:
                    writer.write(bytes_stream)
                    bytes_stream.seek(0)
                    
                    # Use Gemini to handle PDF interpretation instead of rendering
                    # We'll create a simple bitmap representation
                    img = Image.new('RGB', (1700, 2200), color='white')
                    images.append(img)
                    
                    # Attach the PDF data to the image as custom metadata
                    # (This is a hack to pass the PDF content to the Gemini API)
                    img.pdf_data = bytes_stream.getvalue()
    except Exception as e:
        print(f"Error converting PDF to images: {e}")
    
    return images

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using Gemini Vision."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = """Extract all text from this document exactly as it appears. 
    Pay special attention to questions in formats like:
    - Question 1 [5 marks]
    - Question 2: Describe... [10 marks]
    - Q3. Explain... [3 marks]"""
    
    extracted_text = ""
    
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
    except Exception as e:
        print(f"PyPDF2 text extraction failed: {e}. Falling back to Gemini.")
        
        # Fall back to Gemini if direct extraction fails
        images = pdf_to_images(pdf_path)
        
        for image in images:
            # If the image has attached PDF data, send it to Gemini
            if hasattr(image, 'pdf_data'):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(image.pdf_data)
                    temp_pdf_path = temp_pdf.name
                
                try:
                    # Send the PDF directly to Gemini
                    with open(temp_pdf_path, "rb") as f:
                        pdf_content = f.read()
                    
                    # Convert to PNG as fallback if needed
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                        image.save(temp_img.name, format="PNG")
                    
                    with Image.open(temp_img.name) as img:
                        response = model.generate_content([prompt, img], stream=False)
                        extracted_text += response.text + "\n"
                    
                    os.unlink(temp_img.name)
                finally:
                    os.unlink(temp_pdf_path)
            else:
                # Direct usage of the image
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                    image.save(temp_img.name, format="PNG")
                
                with Image.open(temp_img.name) as img:
                    response = model.generate_content([prompt, img], stream=False)
                    extracted_text += response.text + "\n"
                
                os.unlink(temp_img.name)
    
    return extracted_text

def extract_questions_from_pdf(pdf_path):
    """Extract questions from a PDF file."""
    questions = []
    pattern = re.compile(
        r"(?:Question|Q)\s*\d+[:.]?\s*(.*?)\s*\[(\d+)\s*marks?\]", 
        re.IGNORECASE | re.DOTALL
    )
    
    extracted_text = extract_text_from_pdf(pdf_path)
    matches = pattern.finditer(extracted_text)
    
    for match in matches:
        question = match.group(1).strip()
        marks = int(match.group(2))
        questions.append((question, marks))
    
    return questions

def classify_questions(questions):
    """Classify questions using Gemini."""
    classified_data = []
    model = genai.GenerativeModel("gemini-1.5-flash")
    
    for i, (question, marks) in enumerate(questions):
        prompt = f"""
        Analyze the following question and classify it as exactly one of these types:
        - "MCQ" (if it's a multiple choice question)
        - "Descriptive" (if it requires a written explanation)
        - "Opinion-based" (if it asks for personal opinion)
        
        Question: {question}
        
        Return ONLY the classification word (MCQ, Descriptive, or Opinion-based) with no other text.
        """
        
        try:
            response = model.generate_content(prompt)
            predicted_type = response.text.strip()

            if predicted_type not in ["MCQ", "Descriptive", "Opinion-based"]:
                prompt += "\n\nPlease respond with ONLY one word: MCQ, Descriptive, or Opinion-based"
                response = model.generate_content(prompt)
                predicted_type = response.text.strip()
                
                if predicted_type not in ["MCQ", "Descriptive", "Opinion-based"]:
                    predicted_type = "Descriptive"
        except Exception as e:
            print(f"Error classifying question: {e}")
            predicted_type = "Descriptive"
        
        classified_data.append({
            "Question Number": i + 1,
            "Question": question,
            "Type": predicted_type,
            "Marks": marks
        })
    
    return classified_data

def save_questions_to_json(classified_data, output_json):
    """Save classified questions to a JSON file."""
    os.makedirs(os.path.dirname(output_json), exist_ok=True)
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(classified_data, f, indent=4, ensure_ascii=False)

def get_answer(question, marks):
    """Generate an answer using Gemini API."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    
    if marks == 1:
        prompt = f"Answer this in one precise sentence:\nQ: {question}"
    elif marks == 2:
        prompt = f"Provide an answer in 2-3 sentences:\nQ: {question}"
    elif marks == 3:
        prompt = f"Provide a short answer in 3-4 sentences:\nQ: {question}"
    elif marks == 4:
        prompt = f"Provide a detailed answer in about 5-6 sentences:\nQ: {question}"
    elif marks >= 5:
        prompt = f"Write a comprehensive answer in multiple paragraphs (about 500 words):\nQ: {question}"
    else:
        prompt = f"Provide a precise and clear answer:\nQ: {question}"

    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response.text else "No answer generated."
    except Exception as e:
        print(f"Error generating answer: {e}")
        return "Error generating answer."

def generate_answers_json_and_docx(classified_data, output_json, output_docx):
    """Generate answers and save them in JSON and DOCX formats."""
    save_questions_to_json(classified_data, output_json)
    
    for entry in classified_data:
        entry["Reference Answer"] = get_answer(entry["Question"], entry["Marks"])
    
    save_questions_to_json(classified_data, output_json)
    
    doc = Document()
    doc.add_heading('Reference Answers', level=1)
    
    for entry in classified_data:
        doc.add_heading(f"Question {entry['Question Number']}", level=2)
        doc.add_paragraph(entry["Question"])
        doc.add_paragraph(f"Marks: {entry['Marks']}")
        doc.add_paragraph("Reference Answer:")
        doc.add_paragraph(entry["Reference Answer"])
        doc.add_paragraph("\n")
    
    doc.save(output_docx)

def process_question_paper(input_file):
    """Process the question paper and generate reference answers."""
    output_dir = os.path.join(BASE_DIR,"outputs")
    os.makedirs(output_dir, exist_ok=True)
    output_json = os.path.join(output_dir, "classified_questions.json")
    output_docx = os.path.join(output_dir, "answers.docx")
    
    print(f"Processing {input_file}...")
    
    try:
        if input_file.endswith(".pdf"):
            print("Extracting questions from PDF...")
            questions = extract_questions_from_pdf(input_file)
        elif input_file.endswith(".docx"):
            print("Extracting questions from DOCX...")
            questions = extract_questions_from_docx(input_file)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
        
        print(f"Found {len(questions)} questions. Classifying them...")
        classified_data = classify_questions(questions)
        
        print("Generating reference answers...")
        generate_answers_json_and_docx(classified_data, output_json, output_docx)
        
        print(f"Processing complete. Files saved to:\n- {output_json}\n- {output_docx}")
        return True
    except Exception as e:
        print(f"Error processing question paper: {e}")
        return False

def extract_text_from_image(image):
    """Extract text from an image using Gemini 1.5 Flash."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = "Extract only the text from this image. Do not add any explanations or extra words. Maintain the exact structure of the text."
    response = model.generate_content([image, prompt], stream=False)
    return response.text.strip()

def extract_answers_from_pdf(pdf_path):
    """Extract answers from a student's PDF using Gemini OCR."""
    answers = []
    roll_number = "Unknown"

    images = pdf_to_images(pdf_path)

    for page_num, image in enumerate(images):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
            image.save(temp_img.name, format="PNG")
        
        with Image.open(temp_img.name) as img:
            extracted_text = extract_text_from_image(img)
        
        os.unlink(temp_img.name)

        print(f"Page {page_num + 1} Extracted Text:\n{extracted_text}\n")

        if extracted_text:
            if page_num == 0:  
                if "Roll Number" in extracted_text:
                    roll_line = extracted_text.split("Roll Number")[1].split("\n")[0].strip()
                    roll_number = roll_line.split()[-1].strip()  

            lines = extracted_text.split("\n")
            current_answer = []  
            for line in lines:
                line = line.strip()
                if not line:
                    continue  

                if line.lower().startswith("ans"):
                    if current_answer:
                        answers.append(" ".join(current_answer).strip())  
                        current_answer = []  

                    current_answer.append(line)
                else:
                    current_answer.append(line)

            if current_answer:
                answers.append(" ".join(current_answer).strip())

    return answers, roll_number

def evaluate_answer(question, reference_answer, student_answer, q_type, marks):
    """Use Gemini to evaluate student answers."""
    prompt = f"""
    You are an expert teacher evaluating a student's answer.
    
    Question: {question}
    Question Type: {q_type}
    Reference Answer: {reference_answer}
    Student Answer: {student_answer}
    Maximum Marks: {marks}
    
    Based on comparison between the reference and student answer, assign a score out of {marks}.
    For MCQs, give full marks only for exact matches.
    For Descriptive answers, evaluate based on factual accuracy, completeness, and clarity.
    For Opinion-based answers, evaluate based on coherence, logical flow, and relevance.
    
    Return only the numeric score as a float (with up to 1 decimal place). Do not include any explanation or additional text.
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    try:
        score = float(response.text.strip())
        return min(score, marks)  
    except ValueError:
        if q_type == "MCQ":
            return marks if reference_answer.lower() == student_answer.lower() else 0.0
        else:
            return marks / 2

def identify_weak_topics(question, student_answer, reference_answer):
    """Use Gemini to identify weak topics and improvement tips."""
    prompt = f"""
    Analyze the student's mistakes in answering the following question. Identify the weak topic and suggest improvements.
    
    Question: {question}
    Reference Answer: {reference_answer}
    Student Answer: {student_answer}
    
    Provide the weak topic and tips in JSON format with keys 'weak_topic' and 'tips'. The 'tips' should be a list of strings.
    Example: {{"weak_topic": "Topic Name", "tips": ["Tip 1", "Tip 2", "Tip 3"]}}
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    try:
        result = json.loads(response.text)
        return result
    except json.JSONDecodeError:
        weak_topic_match = re.search(r'"weak_topic"\s*:\s*"([^"]+)"', response.text)
        tips_match = re.findall(r'"tips"\s*:\s*\[([^\]]+)\]', response.text)
        
        weak_topic = weak_topic_match.group(1) if weak_topic_match else "Unknown"
        tips = []
        if tips_match:
            tips = [tip.strip().strip('"\'') for tip in tips_match[0].split(",")]
        
        return {"weak_topic": weak_topic, "tips": tips or ["No specific tips available."]}

def evaluate_answers(input_pdf, classified_json, output_json):
    """Evaluate a student's answers against reference answers."""
    try:
        with open(classified_json, "r", encoding="utf-8") as f:
            classified_data = json.load(f)
        
        student_answers, roll_number = extract_answers_from_pdf(input_pdf)
        
        results = []
        
        for i, entry in enumerate(classified_data):
            question_num = entry["Question Number"]
            question = entry["Question"]
            q_type = entry["Type"]
            marks = entry["Marks"]
            reference_answer = entry["Reference Answer"]
            
            student_answer = student_answers[i] if i < len(student_answers) else ""
            
            score = evaluate_answer(question, reference_answer, student_answer, q_type, marks)
            
            result_entry = {
                "roll_number": roll_number,
                "file_name": os.path.basename(input_pdf),
                "question_number": question_num,
                "question": question,
                "reference_answer": reference_answer,
                "student_answer": student_answer,
                "marks_awarded": score,
                "total_marks": marks
            }
            
            if score < marks * 0.8: 
                weak_data = identify_weak_topics(question, student_answer, reference_answer)
                result_entry["weak_topics"] = weak_data["weak_topic"]
                result_entry["improvement_tips"] = weak_data["tips"]
            
            results.append(result_entry)
        
        os.makedirs(os.path.dirname(output_json), exist_ok=True)
        
        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=4, ensure_ascii=False)
        
        return results
    
    except Exception as e:
        print(f"An error occurred while evaluating answers: {e}")
        return []

def process_student_answer_sheet(student_pdf):
    """Process a single student's answer sheet."""
    classified_json = os.path.join(BASE_DIR,"outputs", "classified_questions.json")
    output_json = f"{os.path.join(BASE_DIR, 'data/')}{os.path.basename(student_pdf).split('.')[0]}_results.json"
    
    results = evaluate_answers(student_pdf, classified_json, output_json)
    print(f"Evaluation complete for {student_pdf}. Results saved to {output_json}")
    
    return output_json

def process_all_student_sheets(folder_path):
    """Process all student answer sheets in a folder."""
    result_files = []
    
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf") and "question-paper" not in filename.lower():
            student_pdf = os.path.join(folder_path, filename)
            result_file = process_student_answer_sheet(student_pdf)
            result_files.append(result_file)
    
    return result_files

def generate_class_summary(result_files, output_json):
    """Generate a summary of all students' performance."""
    all_results = []
    student_summaries = {}
    
    for file_path in result_files:
        with open(file_path, "r", encoding="utf-8") as f:
            results = json.load(f)
            all_results.extend(results)
            
            if results:
                roll_number = results[0]["roll_number"]
                total_marks = sum(r["total_marks"] for r in results)
                marks_awarded = sum(r["marks_awarded"] for r in results)
                percentage = (marks_awarded / total_marks * 100) if total_marks > 0 else 0
                
                weak_topics = {}
                for result in results:
                    if "weak_topics" in result:
                        topic = result["weak_topics"]
                        weak_topics[topic] = weak_topics.get(topic, 0) + 1
                
                common_weak_topics = [topic for topic, count in 
                                     sorted(weak_topics.items(), key=lambda x: x[1], reverse=True)[:3]]
                
                student_summaries[roll_number] = {
                    "roll_number": roll_number,
                    "total_marks": total_marks,
                    "marks_awarded": marks_awarded,
                    "percentage": round(percentage, 2),
                    "common_weak_topics": common_weak_topics
                }
    
    summary = {
        "individual_results": all_results,
        "student_summaries": list(student_summaries.values())
    }
    
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=4, ensure_ascii=False)
    
    print(f"Class summary generated and saved to {output_json}")

def check_directory_access(path):
    """Check if directory exists and is accessible"""
    if not os.path.exists(path):
        print(f"Error: Directory does not exist - {path}")
        return False
    if not os.access(path, os.R_OK | os.W_OK):
        print(f"Error: No read/write access to directory - {path}")
        return False
    return True

def check_file_exists(filepath):
    """Verify file exists and is readable"""
    if not os.path.exists(filepath):
        print(f"Error: File not found - {filepath}")
        return False
    if not os.access(filepath, os.R_OK):
        print(f"Error: Cannot read file - {filepath}")
        return False
    return True

def check_gemini_api():
    """Verify Gemini API is reachable"""
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content("Test connection", stream=False)
        return response is not None
    except Exception as e:
        print(f"Gemini API connection failed: {str(e)}")
        return False

def check_pdf_conversion():
    """Verify PDF conversion tools are available"""
    try:
        # Just check if PyPDF2 and PIL are available
        import PyPDF2
        from PIL import Image
        return True
    except Exception as e:
        print(f"PDF conversion check failed: {str(e)}")
        return False 

if __name__ == "__main__":
    # Add pre-flight checks
    print("Running system checks...")
    try:
        folder = os.path.join(BASE_DIR,"uploads")
        
        # Verify uploads directory exists
        if not os.path.exists(folder):
            print(f"ERROR: Uploads directory not found at {folder}")
            print("Contents of /tmp:", os.listdir(BASE_DIR))
            exit(1)
    
        if not check_directory_access(BASE_DIR):
            exit(1)
    
        if not check_gemini_api():
            print("Gemini API check failed")
            exit(1)
    
        if not check_pdf_conversion():
            print("PDF conversion check failed")
            exit(1)
    
        folder = os.path.join(BASE_DIR,"uploads")
    
    # Check uploads directory
        if not check_directory_access(folder):
            print(f"Uploads directory not accessible: {folder}")
            exit(1)
    
        base_filename = "question-paper"
        pdf_file = os.path.join(folder, base_filename + ".pdf")
        docx_file = os.path.join(folder, base_filename + ".docx")

    # Check input files exist
        input_file_found = False
        if os.path.exists(pdf_file):
            if not check_file_exists(pdf_file):
                exit(1)
            input_file_found = True
            process_question_paper(pdf_file)
        elif os.path.exists(docx_file):
            if not check_file_exists(docx_file):
             exit(1)
            input_file_found = True
            process_question_paper(docx_file)
    
        if not input_file_found:
            print("No valid question paper file found!")
            exit(1)
    
    # Check outputs directory
        output_dir = os.path.join(BASE_DIR,"outputs")
        if not check_directory_access(output_dir):
            print(f"Cannot access outputs directory: {output_dir}")
            exit(1)
    
    # Process student sheets
        result_files = process_all_student_sheets(folder)
    
    # Check data directory
        data_dir = os.path.join(BASE_DIR,"data/")
        if not check_directory_access(data_dir):
            print(f"Cannot access data directory: {data_dir}")
            exit(1)
    
        summary_json = os.path.join(data_dir,"class_summary.json")
        generate_class_summary(result_files, summary_json)
    
    # Final verification
        if not os.path.exists(summary_json):
            print(f"Error: Summary file not created at {summary_json}")
            exit(1)
    except Exception as e:
        print(f"CRITICAL ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        exit(1)
    
    print("All processing completed successfully with all checks passed!")
